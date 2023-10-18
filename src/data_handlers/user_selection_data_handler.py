from docx import Document
from typing import Dict, List
from .base_data_handler import BaseDataHandler


class UserSelectionDataHandler(BaseDataHandler):
    # def __init__(self):
    # super().__init()

    def get_selection(self) -> Dict[str, Dict[str, List[str]]]:
        """
        Get the selected sentences.

        :return: A dictionary containing selected sentences and words, categorized by sections.
        """
        data = {"sentences": self.sentences, "words": self.words}
        return data

    def write_selection_to_docx(self, output_file_path: str):
        """
        Write all selected sentences and words to a .docx document.

        :param output_file_path: The path where the .docx document will be saved.
        """
        document = Document()

        document.add_heading("Sentences", level=1)
        # Iterate through sections and sentences
        for section, sentences in self.sentences.items():
            if sentences:
                document.add_heading(section, level=2)
                for sentence in sentences:
                    document.add_paragraph(sentence)

        #document.add_heading("Words", level=1)
        # Iterate through sections and words
        for section, words in self.words.items():
            if words:
                document.add_heading(section, level=2)
                for word in words:
                    document.add_paragraph(word)

        # Save the document to the specified output file path
        try:
            document.save(output_file_path)
            print(f"Document saved to {output_file_path}")
        except Exception as e:
            print(f"An error occurred while saving the document: {str(e)}")
